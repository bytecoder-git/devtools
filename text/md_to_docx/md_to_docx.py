#!/usr/bin/env python3
import sys
import argparse
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from bs4 import BeautifulSoup

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)
    run._element.append(instrText)
    run._element.append(fldChar2)

def convert_table_to_docx(table_html, document):
    soup = BeautifulSoup(table_html, 'html.parser')
    table = soup.find('table')

    if table:
        rows = table.find_all('tr')
        docx_table = document.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
        docx_table.style = 'Table Grid'

        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                docx_table.cell(i, j).text = cell.get_text(strip=True)

        document.add_paragraph()  # Add a blank line after the table

def extract_title_from_markdown(markdown_content):
    for line in markdown_content.split('\n'):
        if line.strip().startswith('# '):
            return line.strip()[2:]  # Return the text after '# '
    return "Untitled"  # Return "Untitled" if no Header 1 is found

def add_section_break(document):
    new_section = document.add_section(WD_SECTION.NEW_PAGE)
    sectPr = new_section._sectPr
    type_element = sectPr.xpath('./w:type')
    if type_element:
        type_element[0].set(qn('w:val'), 'nextPage')
    else:
        new_type = OxmlElement('w:type')
        new_type.set(qn('w:val'), 'nextPage')
        sectPr.append(new_type)
    return new_section

def convert_markdown_to_docx(markdown_content, output_file, book_title, date_str):
    html_content = markdown.markdown(markdown_content, extensions=['tables'])
    document = Document()

    # Add title page
    document.add_heading(book_title, 0)
    title_section = document.sections[0]
    title_section.start_page_number = 0  # Exclude title page from numbering

    # Add a section break after title page
    content_section = add_section_break(document)
    content_section.start_page_number = 1  # Start page numbering from 1 for content

    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')

    # Process content
    current_section = None

    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'table', 'hr', 'ul', 'ol']):
        if element.name == 'hr':
            add_section_break(document)
        elif element.name.startswith('h'):
            level = int(element.name[1])
            if level == 1:
                current_section = element.get_text()
            document.add_heading(element.get_text(), level=level)
        elif element.name == 'table':
            convert_table_to_docx(str(element), document)
        elif element.name == 'p':
            paragraph = document.add_paragraph(element.get_text())
            paragraph.style = document.styles['Normal']
        elif element.name in ['ul', 'ol']:
            # Start a new list
            list_style = 'List Number' if element.name == 'ol' else 'List Bullet'
            for li in element.find_all('li'):
                paragraph = document.add_paragraph(li.get_text(), style=list_style)
                if element.name == 'ol':
                    # Reset numbering for each new ordered list
                    paragraph._element.get_or_add_pPr().get_or_add_numPr().get_or_add_numId().val = 1

    # Add headers and footers, excluding the title page
    for i, section in enumerate(document.sections):
        # Set up header
        header = section.header
        header.is_linked_to_previous = False
        if i > 0:  # Skip the title page
            header_para = header.paragraphs[0]
            header_para.text = book_title
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set up footer
        footer = section.footer
        footer.is_linked_to_previous = False
        if i > 0:  # Skip the title page
            for paragraph in footer.paragraphs:
                paragraph.clear()

            footer_para = footer.paragraphs[0]

            # Add date to the left
            date_run = footer_para.add_run(date_str)
            date_run.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add section name to the center
            section_name = current_section if current_section else f"Section {i}"
            footer_para.add_run('\t' + section_name + '\t')

            # Add "Page {#}" to the right
            page_run = footer_para.add_run("Page ")
            add_page_number(footer_para)

            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Save the document
    document.save(output_file)

def main():
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX", formatter_class=argparse.RawTextHelpFormatter)
    parser.add_argument('-i', '--input', help='Input Markdown file')
    parser.add_argument('-o', '--output', help='Output DOCX file (default: output.docx)')
    parser.add_argument('-t', '--title', help='Book title (default: extracted from first H1 in Markdown)')
    parser.add_argument('-d', '--date', help='Date to include in footer (format: YYYY-MM-DD, default: current date)')

    # Check if there's input on stdin or any arguments
    if sys.stdin.isatty() and len(sys.argv) == 1:
        print("Usage: python md_to_docx.py [OPTIONS]\n")
        print("Options:")
        print("  -h, --help            Show this help message and exit")
        print("  -i, --input INPUT     Input Markdown file")
        print("  -o, --output OUTPUT   Output DOCX file (default: output.docx)")
        print("  -t, --title TITLE     Book title (default: extracted from first H1 in Markdown)")
        print("  -d, --date DATE       Date to include in footer (format: YYYY-MM-DD, default: current date)")
        print("\nExamples:")
        print("  python md_to_docx.py -i input.md -o output.docx")
        print("  python md_to_docx.py -i input.md -t 'My Book' -d 2024-08-28")
        print("  cat input.md | python md_to_docx.py -o output.docx")
        sys.exit(1)

    args = parser.parse_args()

    # Read markdown content
    if args.input:
        with open(args.input, 'r', encoding='utf-8') as file:
            markdown_content = file.read()
    elif not sys.stdin.isatty():
        markdown_content = sys.stdin.read()
    else:
        print("Error: No input provided. Please specify an input file or pipe content to stdin.")
        parser.print_help()
        sys.exit(1)

    # Determine book title
    if args.title:
        book_title = args.title
    else:
        book_title = extract_title_from_markdown(markdown_content)

    # Determine output file name
    if args.output:
        output_file = args.output if args.output.endswith('.docx') else args.output + '.docx'
    else:
        output_file = 'output.docx'

    # Determine date
    if args.date:
        try:
            date_obj = datetime.strptime(args.date, "%Y-%m-%D")
            date_str = date_obj.strftime("%B %d, %Y")
        except ValueError:
            print("Invalid date format. Using current date.")
            date_str = datetime.now().strftime("%B %d, %Y")
    else:
        date_str = datetime.now().strftime("%B %d, %Y")

    # Convert markdown to docx
    convert_markdown_to_docx(markdown_content, output_file, book_title, date_str)
    print(f"Conversion complete. Output saved to {output_file}")

if __name__ == "__main__":
    main()

if __name__ == "__main__":
    main()
